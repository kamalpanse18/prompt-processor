
# Multi-Modal Input Processor
# Handles text, images, PDFs, and DOCX files


import os
import io
from typing import List, Dict, Optional
from pathlib import Path

# For PDF processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# For DOCX processing
try:
    from docx import Document
except ImportError:
    Document = None

# For image processing (OCR)
try:
    from PIL import Image
    import pytesseract
    
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

    IMAGE_AVAILABLE = True
except ImportError:
    Image = None
    pytesseract = None

class InputProcessor:
    """Processes various input types and extracts text content"""
    
    def __init__(self):
        self.supported_formats = {
            'text': ['.txt', '.md'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.gif'],
            'pdf': ['.pdf'],
            'docx': ['.docx']
        }
    
    def process_input(self, input_path: str) -> Dict:
        """Main entry point for processing any input"""
        if not os.path.exists(input_path):
            if not any(char in input_path for char in ['.', '/', '\\']):
                return self.process_text(input_path)
            else:
                return {
                    'type': 'error',
                    'content': '',
                    'metadata': {},
                    'success': False,
                    'error': 'File not found'
                }
        
        file_ext = Path(input_path).suffix.lower()
        
        if file_ext in self.supported_formats['text']:
            return self.process_text_file(input_path)
        elif file_ext in self.supported_formats['image']:
            return self.process_image(input_path)
        elif file_ext in self.supported_formats['pdf']:
            return self.process_pdf(input_path)
        elif file_ext in self.supported_formats['docx']:
            return self.process_docx(input_path)
        else:
            return {
                'type': 'error',
                'content': '',
                'metadata': {},
                'success': False,
                'error': f'Unsupported file format: {file_ext}'
            }
    
    def process_text(self, text: str) -> Dict:
        """Process plain text input"""
        return {
            'type': 'text',
            'content': text.strip(),
            'metadata': {
                'length': len(text),
                'source': 'direct_input'
            },
            'success': True
        }
    
    def process_text_file(self, file_path: str) -> Dict:
        """Process text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                'type': 'text',
                'content': content.strip(),
                'metadata': {
                    'filename': os.path.basename(file_path),
                    'length': len(content),
                    'source': 'file'
                },
                'success': True
            }
        except Exception as e:
            return {
                'type': 'error',
                'content': '',
                'metadata': {},
                'success': False,
                'error': str(e)
            }
    
    def process_image(self, image_path: str) -> Dict:
        """Process image using OCR"""
        if Image is None or pytesseract is None:
            return {
                'type': 'error',
                'content': '',
                'metadata': {},
                'success': False,
                'error': 'PIL or pytesseract not installed'
            }
        
        try:
            img = Image.open(image_path)
            text = pytesseract.image_to_string(img)
            
            return {
                'type': 'image',
                'content': text.strip(),
                'metadata': {
                    'filename': os.path.basename(image_path),
                    'size': img.size,
                    'format': img.format,
                    'source': 'ocr'
                },
                'success': True
            }
        except Exception as e:
            return {
                'type': 'error',
                'content': '',
                'metadata': {},
                'success': False,
                'error': f'Image processing error: {str(e)}'
            }
    
    def process_pdf(self, pdf_path: str) -> Dict:
        """Process PDF file"""
        if PyPDF2 is None:
            return {
                'type': 'error',
                'content': '',
                'metadata': {},
                'success': False,
                'error': 'PyPDF2 not installed'
            }
        
        try:
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return {
                'type': 'pdf',
                'content': text.strip(),
                'metadata': {
                    'filename': os.path.basename(pdf_path),
                    'num_pages': num_pages,
                    'length': len(text),
                    'source': 'pdf_extraction'
                },
                'success': True
            }
        except Exception as e:
            return {
                'type': 'error',
                'content': '',
                'metadata': {},
                'success': False,
                'error': f'PDF processing error: {str(e)}'
            }
    
    def process_docx(self, docx_path: str) -> Dict:
        """Process DOCX file"""
        if Document is None:
            return {
                'type': 'error',
                'content': '',
                'metadata': {},
                'success': False,
                'error': 'python-docx not installed'
            }
        
        try:
            doc = Document(docx_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_text.append(row_text)
            
            if table_text:
                text += "\n\nTables:\n" + "\n".join(table_text)
            
            return {
                'type': 'docx',
                'content': text.strip(),
                'metadata': {
                    'filename': os.path.basename(docx_path),
                    'num_paragraphs': len(doc.paragraphs),
                    'num_tables': len(doc.tables),
                    'length': len(text),
                    'source': 'docx_extraction'
                },
                'success': True
            }
        except Exception as e:
            return {
                'type': 'error',
                'content': '',
                'metadata': {},
                'success': False,
                'error': f'DOCX processing error: {str(e)}'
            }
    
    def process_multiple_inputs(self, input_paths: List[str]) -> Dict:
        """Process multiple inputs and combine them"""
        results = []
        combined_content = []
        combined_metadata = {
            'sources': [],
            'types': []
        }
        
        for input_path in input_paths:
            result = self.process_input(input_path)
            results.append(result)
            
            if result['success']:
                combined_content.append(f"--- Source: {result.get('metadata', {}).get('filename', 'text input')} ---")
                combined_content.append(result['content'])
                combined_metadata['sources'].append(result.get('metadata', {}).get('filename', 'text'))
                combined_metadata['types'].append(result['type'])
        
        return {
            'type': 'mixed',
            'content': "\n\n".join(combined_content),
            'metadata': combined_metadata,
            'individual_results': results,
            'success': all(r['success'] for r in results)
        }